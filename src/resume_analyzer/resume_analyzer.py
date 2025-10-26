"""
Resume Analyzer - Compares job descriptions with resumes using OpenAI API
Provides similarity scores, matching/missing skills, and improvement recommendations
"""

import os
import json
import re
from typing import Dict, List, Tuple
from pathlib import Path
from openai import OpenAI
from dotenv import load_dotenv

# Document processing libraries
import PyPDF2
from docx import Document
import pdfplumber

# Load environment variables
load_dotenv()


class ResumeAnalyzer:
    """Analyzes resumes against job descriptions using OpenAI's GPT models."""

    def __init__(self, api_key: str = None):
        """
        Initialize the analyzer with OpenAI API key.

        Args:
            api_key: OpenAI API key (optional if set in environment)
        """
        self.api_key = api_key or os.getenv('OPENAI_API_KEY')
        if not self.api_key:
            raise ValueError("OpenAI API key must be provided or set in OPENAI_API_KEY environment variable")

        self.client = OpenAI(api_key=self.api_key)
        self.model = "gpt-4o-mini"  # Latest stable model, cost-effective

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from various file formats (txt, pdf, docx).

        Args:
            file_path: Path to the resume or job description file

        Returns:
            Extracted text content

        Supported formats: .txt, .pdf, .docx
        Note: .doc files are not supported. Please convert to .docx format.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Get file extension
        file_extension = Path(file_path).suffix.lower()

        # Check for unsupported .doc format
        if file_extension == '.doc':
            raise Exception(
                "Legacy .doc format is not supported. "
                "Please convert your file to .docx format using Microsoft Word or a free converter."
            )

        try:
            # PDF files
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)

            # DOCX files
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)

            # Text files
            elif file_extension in ['.txt', '.text']:
                return self._extract_from_txt(file_path)

            else:
                # Try reading as text file by default
                return self._extract_from_txt(file_path)

        except Exception as e:
            raise Exception(f"Error reading file {file_path}: {str(e)}")

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files using multiple methods for better accuracy."""
        text = ""

        # Method 1: Try pdfplumber first (better for complex PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            if text.strip():
                return text
        except Exception as e:
            print(f"Warning: pdfplumber failed, trying PyPDF2: {str(e)}")

        # Method 2: Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"

            if text.strip():
                return text
            else:
                raise Exception("No text could be extracted from PDF")
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files."""
        try:
            doc = Document(file_path)
            text = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return "\n".join(text)
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from DOC files (older Word format)."""
        import sys
        import subprocess

        # Method 1: Try using pywin32 on Windows
        if sys.platform == 'win32':
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(file_path))
                text = doc.Content.Text
                doc.Close()
                word.Quit()
                return text
            except ImportError:
                pass  # pywin32 not installed, try other methods
            except Exception as e:
                print(f"Warning: Failed to extract using pywin32: {str(e)}")

        # Method 2: Try using antiword (Linux/Mac command-line tool)
        try:
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                check=True
            )
            return result.stdout
        except FileNotFoundError:
            pass  # antiword not installed
        except subprocess.CalledProcessError as e:
            print(f"Warning: antiword failed: {str(e)}")
        except Exception as e:
            print(f"Warning: antiword error: {str(e)}")

        # Method 3: Try using LibreOffice in headless mode (cross-platform)
        try:
            import tempfile
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert to txt using LibreOffice
                subprocess.run(
                    ['soffice', '--headless', '--convert-to', 'txt:Text',
                     '--outdir', temp_dir, file_path],
                    capture_output=True,
                    check=True,
                    timeout=30
                )

                # Read the converted file
                txt_file = os.path.join(temp_dir, Path(file_path).stem + '.txt')
                if os.path.exists(txt_file):
                    with open(txt_file, 'r', encoding='utf-8') as f:
                        return f.read()
        except FileNotFoundError:
            pass  # LibreOffice not installed
        except subprocess.CalledProcessError:
            pass  # Conversion failed
        except Exception as e:
            print(f"Warning: LibreOffice conversion failed: {str(e)}")

        # Method 4: Try using pypandoc
        try:
            import pypandoc
            text = pypandoc.convert_file(file_path, 'plain', format='doc')
            return text
        except ImportError:
            pass  # pypandoc not installed
        except Exception as e:
            print(f"Warning: pypandoc failed: {str(e)}")

        # If all methods fail, provide helpful error message
        raise Exception(
            "Failed to extract text from .doc file. Please try one of these solutions:\n"
            "1. Convert the .doc file to .docx format (recommended)\n"
            "2. On Windows: Install pywin32 with 'pip install pywin32'\n"
            "3. On Linux/Mac: Install antiword with 'sudo apt-get install antiword' or 'brew install antiword'\n"
            "4. Install LibreOffice (soffice command) for cross-platform support\n"
            "5. Install pypandoc with 'pip install pypandoc'\n"
            "\nNote: .docx files are fully supported without additional dependencies."
        )

    def calculate_similarity_score(self, job_description: str, resume: str) -> Dict:
        """
        Calculate similarity score between job description and resume.

        Args:
            job_description: Job description text
            resume: Resume text

        Returns:
            Dictionary containing similarity analysis
        """
        prompt = f"""
You are an expert ATS (Applicant Tracking System) and recruitment specialist.

Analyze the following job description and resume, then provide a detailed similarity score and analysis.

JOB DESCRIPTION:
{job_description}

RESUME:
{resume}

Please provide your analysis in the following JSON format:
{{
    "similarity_score": <number between 0-100>,
    "overall_match": "<Poor/Fair/Good/Excellent>",
    "key_strengths": ["strength1", "strength2", "strength3"],
    "analysis_summary": "Brief summary of how well the resume matches the job"
}}

Be objective and thorough in your assessment.
"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert ATS system and recruitment specialist. Provide accurate, objective analysis in valid JSON format."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                response_format={"type": "json_object"}
            )

            result = json.loads(response.choices[0].message.content)
            return result

        except Exception as e:
            raise Exception(f"Error calculating similarity score: {str(e)}")

    def extract_skills(self, job_description: str, resume: str) -> Dict[str, List[str]]:
        """
        Extract matching and missing skills.

        Args:
            job_description: Job description text
            resume: Resume text

        Returns:
            Dictionary with matching_skills and missing_skills lists
        """
        prompt = f"""
You are an expert technical recruiter and skills analyst.

Analyze the following job description and resume to identify skills.

JOB DESCRIPTION:
{job_description}

RESUME:
{resume}

Identify:
1. Technical skills, tools, and technologies mentioned in the job description
2. Which of these skills are present in the resume (matching skills)
3. Which required/preferred skills are missing from the resume (missing skills)

Provide your analysis in the following JSON format:
{{
    "matching_skills": ["skill1", "skill2", "skill3"],
    "missing_skills": ["skill1", "skill2", "skill3"],
    "partial_matches": ["skill1: explanation", "skill2: explanation"]
}}

Be specific and list actual skill names (e.g., "Python", "AWS", "Agile", "Leadership").
"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert technical recruiter. Provide accurate skills analysis in valid JSON format."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                response_format={"type": "json_object"}
            )

            result = json.loads(response.choices[0].message.content)
            return result

        except Exception as e:
            raise Exception(f"Error extracting skills: {str(e)}")

    def generate_recommendations(self, job_description: str, resume: str,
                                missing_skills: List[str]) -> List[str]:
        """
        Generate actionable recommendations to improve the resume.

        Args:
            job_description: Job description text
            resume: Resume text
            missing_skills: List of skills missing from resume

        Returns:
            List of recommendations
        """
        prompt = f"""
You are an expert resume writer and career coach.

Based on the job description and current resume, provide specific, actionable recommendations to improve the resume.

JOB DESCRIPTION:
{job_description}

RESUME:
{resume}

MISSING SKILLS:
{', '.join(missing_skills) if missing_skills else 'None identified'}

Provide 5-8 specific, actionable recommendations in the following JSON format:
{{
    "recommendations": [
        "Specific recommendation 1",
        "Specific recommendation 2",
        "Specific recommendation 3"
    ]
}}

Focus on:
- How to highlight relevant experience better
- Keywords to add (if genuinely applicable)
- Format/structure improvements
- Ways to address missing skills
- Quantifying achievements
- Tailoring the resume to the job
"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert resume writer. Provide specific, actionable recommendations in valid JSON format."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.4,
                response_format={"type": "json_object"}
            )

            result = json.loads(response.choices[0].message.content)
            return result.get('recommendations', [])

        except Exception as e:
            raise Exception(f"Error generating recommendations: {str(e)}")

    def analyze(self, job_description: str, resume: str) -> Dict:
        """
        Perform complete analysis of resume against job description.

        Args:
            job_description: Job description text or file path
            resume: Resume text or file path

        Returns:
            Complete analysis results
        """
        # Check if inputs are file paths
        if os.path.isfile(job_description):
            job_description = self.extract_text_from_file(job_description)

        if os.path.isfile(resume):
            resume = self.extract_text_from_file(resume)

        # Validate inputs
        if not job_description.strip():
            raise ValueError("Job description cannot be empty")
        if not resume.strip():
            raise ValueError("Resume cannot be empty")

        print("Analyzing resume against job description...")
        print("-" * 50)

        # Calculate similarity score
        print("⏳ Calculating similarity score...")
        similarity_result = self.calculate_similarity_score(job_description, resume)

        # Extract skills
        print("⏳ Extracting skills...")
        skills_result = self.extract_skills(job_description, resume)

        # Generate recommendations
        print("⏳ Generating recommendations...")
        recommendations = self.generate_recommendations(
            job_description,
            resume,
            skills_result.get('missing_skills', [])
        )

        # Compile results
        results = {
            'similarity_score': similarity_result.get('similarity_score', 0),
            'overall_match': similarity_result.get('overall_match', 'Unknown'),
            'key_strengths': similarity_result.get('key_strengths', []),
            'analysis_summary': similarity_result.get('analysis_summary', ''),
            'matching_skills': skills_result.get('matching_skills', []),
            'missing_skills': skills_result.get('missing_skills', []),
            'partial_matches': skills_result.get('partial_matches', []),
            'recommendations': recommendations
        }

        return results

    def print_results(self, results: Dict) -> None:
        """
        Print analysis results in a formatted way.

        Args:
            results: Analysis results dictionary
        """
        print("\n" + "=" * 70)
        print("RESUME ANALYSIS RESULTS")
        print("=" * 70)

        # Similarity Score
        print(f"\n📊 SIMILARITY SCORE: {results['similarity_score']}/100")
        print(f"   Overall Match: {results['overall_match']}")

        # Analysis Summary
        if results.get('analysis_summary'):
            print(f"\n📝 SUMMARY:")
            print(f"   {results['analysis_summary']}")

        # Key Strengths
        if results.get('key_strengths'):
            print(f"\n✅ KEY STRENGTHS:")
            for i, strength in enumerate(results['key_strengths'], 1):
                print(f"   {i}. {strength}")

        # Matching Skills
        if results.get('matching_skills'):
            print(f"\n✨ MATCHING SKILLS ({len(results['matching_skills'])}):")
            for skill in results['matching_skills']:
                print(f"   • {skill}")

        # Partial Matches
        if results.get('partial_matches'):
            print(f"\n⚠️  PARTIAL MATCHES ({len(results['partial_matches'])}):")
            for match in results['partial_matches']:
                print(f"   • {match}")

        # Missing Skills
        if results.get('missing_skills'):
            print(f"\n❌ MISSING SKILLS ({len(results['missing_skills'])}):")
            for skill in results['missing_skills']:
                print(f"   • {skill}")

        # Recommendations
        if results.get('recommendations'):
            print(f"\n💡 RECOMMENDATIONS ({len(results['recommendations'])}):")
            for i, rec in enumerate(results['recommendations'], 1):
                print(f"   {i}. {rec}")

        print("\n" + "=" * 70)

    def save_results(self, results: Dict, output_file: str = "resume_analysis.json") -> None:
        """
        Save analysis results to a JSON file.

        Args:
            results: Analysis results dictionary
            output_file: Output file path
        """
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            print(f"\n💾 Results saved to: {output_file}")
        except Exception as e:
            print(f"\n⚠️  Error saving results: {str(e)}")


def main():
    """Main function to demonstrate usage."""

    # Example usage
    print("Resume Analyzer with OpenAI")
    print("=" * 70)
    print("\nSupported file formats: .txt, .pdf, .docx")
    print("=" * 70)

    # Initialize analyzer
    try:
        analyzer = ResumeAnalyzer()
    except ValueError as e:
        print(f"Error: {e}")
        print("\nPlease set your OpenAI API key:")
        print("1. Create a .env file with: OPENAI_API_KEY=your-key-here")
        print("2. Or set it as an environment variable")
        print("3. Or pass it directly: ResumeAnalyzer(api_key='your-key')")
        return

    # Example job description and resume
    job_description = """
Senior Software Engineer - Python/AI

We are seeking an experienced Senior Software Engineer with strong Python expertise
and AI/ML knowledge to join our team.

Requirements:
- 5+ years of Python development experience
- Experience with machine learning frameworks (TensorFlow, PyTorch)
- Strong understanding of REST APIs and microservices
- Experience with cloud platforms (AWS, GCP, or Azure)
- Proficiency in SQL and database design
- Experience with Docker and Kubernetes
- Knowledge of CI/CD pipelines
- Strong problem-solving and communication skills

Preferred:
- Experience with LangChain or similar LLM frameworks
- React or frontend development experience
- Agile/Scrum methodology experience
"""

    resume = """
John Doe
Software Engineer

Experience:
Software Engineer at Tech Corp (3 years)
- Developed Python applications for data processing
- Built REST APIs using Flask and FastAPI
- Worked with PostgreSQL databases
- Implemented CI/CD pipelines using Jenkins
- Collaborated with cross-functional teams

Skills:
- Programming: Python, JavaScript, Java
- Frameworks: Flask, FastAPI, Django
- Databases: PostgreSQL, MongoDB
- Tools: Git, Docker, Jenkins
- Cloud: AWS (EC2, S3, Lambda)

Education:
Bachelor's in Computer Science
"""

    # Example with file paths (uncomment to use):
    job_description = "job_description.pdf"  # Can be .txt, .pdf, .docx, .doc
    resume = "resume.pdf"  # Can be .txt, .pdf, .docx, .doc

    # Perform analysis
    try:
        results = analyzer.analyze(job_description, resume)
        analyzer.print_results(results)
        analyzer.save_results(results)

    except Exception as e:
        print(f"\n❌ Error during analysis: {str(e)}")


if __name__ == "__main__":
    main()
